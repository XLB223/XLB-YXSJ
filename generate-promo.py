# -*- coding: utf-8 -*-
"""生成《跨境 AI Listing 生成器》宣传推广资料 Word 文档"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUTPUT = Path(__file__).parent / "跨境AI Listing生成器-宣传推广资料.docx"
SITE_URL = "https://www.kjdsai.cn/"


def set_doc_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def build():
    doc = Document()
    set_doc_font(doc)

    add_title(doc, "跨境 AI Listing 生成器")
    add_subtitle(doc, "宣传推广资料包 · 可直接用于朋友圈 / 小红书 / 抖音 / 卖家群")
    add_subtitle(doc, f"官网：{SITE_URL}")
    doc.add_paragraph()

    add_heading(doc, "一、产品一句话介绍")
    add_para(
        doc,
        "跨境 AI Listing 生成器（kjdsai.cn）是专为 Amazon 跨境卖家打造的在线 AI 工具。"
        "用中文填写产品信息，一键批量生成 36 种语言的 Amazon Listing（标题、五点描述、"
        "HTML 产品描述、后台关键词），支持合规自检与 Word 导出，网页版 + 手机版开箱即用。",
    )
    add_para(doc, "Slogan：输入中文，一键出全球 Amazon Listing", bold=True)
    doc.add_paragraph()

    add_heading(doc, "二、目标用户与核心痛点")
    add_bullets(
        doc,
        [
            "Amazon FBA / FBM 卖家：上新慢、多站点文案写不过来",
            "运营 / 助理：外语能力有限，翻译腔重、不合当地 SEO",
            "代运营 / 外包团队：客户多、SKU 多，人力成本高",
            "多站点扩张团队：欧美日韩中东等站点需要批量本地化 Listing",
            "痛点：ChatGPT 通用对话不专业；Coze 工作流搭建麻烦；人工写 36 语言不现实",
        ],
    )
    doc.add_paragraph()

    add_heading(doc, "三、核心功能亮点")
    add_bullets(
        doc,
        [
            "36 种语言：覆盖亚洲 / 欧洲 / 非洲 / 美洲 Amazon 及主流跨境平台站点",
            "中文输入：产品名、卖点、类目用中文填，AI 按目标语言母语级创作（非直译）",
            "完整 Listing 结构：标题 + 5 条 Bullet Points + HTML 描述 + Search Terms",
            "Amazon 规范内置：类目标题字符上限、249 字节后台词、移动端前 80 字符预览",
            "可选增强：品牌名、核心关键词、竞品参考（学习结构不抄袭）",
            "合规自检：自动标红 Best/Free/医疗宣称等常见敏感词",
            "两种文案风格：专业正式 / 活泼口语化",
            "不编造功能：只基于用户填写的真实卖点生成，避免 AI 乱写",
            "一键导出 Word：多语言结果打包下载，方便团队审核归档",
            "网页版 + 手机版：电脑办公 / 手机随时生成，历史记录自动保存",
        ],
    )
    doc.add_paragraph()

    add_heading(doc, "四、使用流程（3 步上手）")
    add_numbered = lambda items: [
        doc.add_paragraph(item, style="List Number") for item in items
    ]
    for item in [
        f"打开 {SITE_URL}（或手机版 /mobile/）",
        "填写产品名称、卖点、产品类型，勾选目标语言",
        "点击「生成 Listing」→ 查看结果 → 复制或导出 Word",
    ]:
        p = doc.add_paragraph(item, style="List Number")
        for run in p.runs:
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    doc.add_paragraph()

    add_heading(doc, "五、定价方案")
    add_bullets(
        doc,
        [
            "免费试用：每天 3 次，无需注册，打开即用",
            "月卡：¥29.9 / 1 个月，无限次生成",
            "半年卡：¥69.9 / 6 个月（划算）",
            "年卡：¥129.9 / 1 年（最省）",
            "开通方式：官网扫码支付 → 联系客服获取激活码 → 输入激活",
        ],
    )
    doc.add_paragraph()

    add_heading(doc, "六、与常见方案对比")
    add_para(doc, "vs 通用 ChatGPT / DeepSeek 对话", bold=True)
    add_bullets(
        doc,
        [
            "本产品：专为 Amazon Listing 设计，Prompt 内置 SEO / 合规 / 字符限制",
            "通用 AI：需自己写 Prompt，格式不稳定，容易遗漏后台关键词规范",
        ],
    )
    add_para(doc, "vs Coze / 自建 Agent 工作流", bold=True)
    add_bullets(
        doc,
        [
            "Coze：需搭建工作流、配置插件，维护成本高，多语言需反复调试",
            "本产品：打开网页填表即生成，36 语言批量出稿，零代码零维护",
        ],
    )
    add_para(doc, "vs 人工翻译 / 外包", bold=True)
    add_bullets(
        doc,
        [
            "人工：单 SKU 多语言动辄数百上千元，周期长",
            "本产品：年卡 ¥129.9 无限次，几分钟出 36 语言初稿，人工只做终审",
        ],
    )
    doc.add_paragraph()

    add_heading(doc, "七、朋友圈文案（3 条，直接复制）")
    add_para(doc, "【文案 A · 痛点型】", bold=True)
    add_para(
        doc,
        "做 Amazon 跨境最烦什么？写 Listing！\n"
        "36 个站点 36 种语言，标题五点描述后台词……写到头秃。\n"
        f"试试这个工具 👉 {SITE_URL}\n"
        "中文填产品信息，一键出多语言 Amazon Listing\n"
        "免费每天 3 次，先测再决定 🚀",
    )
    add_para(doc, "【文案 B · 功能型】", bold=True)
    add_para(
        doc,
        "分享一个跨境卖家实用工具：跨境 AI Listing 生成器\n"
        "✅ 36 语言批量生成\n"
        "✅ 标题+五点+描述+后台关键词\n"
        "✅ Amazon 合规自检 + Word 导出\n"
        f"👉 {SITE_URL}\n"
        "免费试用 3 次/天，运营党必备",
    )
    add_para(doc, "【文案 C · 效率型】", bold=True)
    add_para(
        doc,
        "以前写一个 SKU 的英语 Listing 要 2 小时\n"
        "现在中文填表，10 分钟出 10 种语言初稿\n"
        "合规检查、字符限制、SEO 结构都帮你做好了\n"
        f"工具地址：{SITE_URL}\n"
        "跨境卖家可以收藏一下",
    )
    doc.add_paragraph()

    add_heading(doc, "八、小红书笔记模板")
    add_para(doc, "标题：跨境卖家必备！中文一键生成 36 语言 Amazon Listing", bold=True)
    add_para(
        doc,
        "正文：\n"
        "做 Amazon 的姐妹有没有同感——\n"
        "上新最耗时的不是找货，是写 Listing！\n"
        "尤其是多站点扩张，每种语言都要重新写，不是翻译就完事，还要符合当地 SEO…\n\n"
        "最近试了一个专门做 Listing 的 AI 工具，记录一下：\n\n"
        "🌐 36 种语言可选（欧美日韩中东都有）\n"
        "📝 中文填产品名+卖点，自动生成标题/五点/描述/后台词\n"
        "✅ 内置 Amazon 规范（标题字符、249 字节后台词）\n"
        "🔍 合规敏感词自动标红\n"
        "📄 可导出 Word 给团队审核\n\n"
        f"地址：{SITE_URL}\n"
        "免费每天 3 次，够测几个 SKU\n\n"
        "#跨境电商 #Amazon #FBA #Listing优化 #亚马逊运营 #AI工具 #跨境工具",
    )
    doc.add_paragraph()

    add_heading(doc, "九、抖音 / 短视频脚本（30 秒）")
    add_para(doc, "【脚本 · 痛点对比】", bold=True)
    add_bullets(
        doc,
        [
            "0-3 秒（钩子）：「做跨境的，写 Listing 还在一个个翻译？」",
            "3-10 秒：屏幕录屏填中文产品信息",
            "10-18 秒：点击生成，展示多语言 Tab 切换结果",
            "18-25 秒：特写合规标红、Word 导出按钮",
            f"25-30 秒（CTA）：「免费试用 3 次/天，搜 kjdsai.cn」",
        ],
    )
    add_para(doc, "【脚本 · 效率展示】", bold=True)
    add_bullets(
        doc,
        [
            "0-5 秒：「10 分钟，36 种语言 Listing，怎么做到的？」",
            "5-15 秒：快放填表 → 生成 → 结果展示",
            "15-25 秒：对比人工 2 小时 vs AI 10 分钟",
            f"25-30 秒：「跨境 AI Listing 生成器，{SITE_URL}」",
        ],
    )
    doc.add_paragraph()

    add_heading(doc, "十、卖家群 / 社群推广话术")
    add_para(
        doc,
        "各位老板好，分享一个我自己在用的 Amazon Listing AI 工具，"
        "专门做跨境 Listing 的，不是通用 ChatGPT。\n\n"
        "特点：中文填信息 → 36 语言批量出标题/五点/描述/后台词，"
        "带 Amazon 合规检查和 Word 导出。\n\n"
        f"免费试用：{SITE_URL}（每天 3 次）\n"
        "有需要多站点上新的可以试一下，觉得好用再开会员。",
    )
    doc.add_paragraph()

    add_heading(doc, "十一、常见问答（FAQ）")
    faq = [
        ("AI 会编造产品功能吗？", "不会。系统 Prompt 明确要求只使用您填写的真实卖点，未提供的信息一律不写。"),
        ("生成的是翻译还是原创？", "不是直译，是按目标语言市场习惯重新创作的 Listing，符合当地 Amazon SEO。"),
        ("支持哪些 Amazon 站点语言？", "36 种，覆盖美国、欧洲五国、日本、中东、东南亚、非洲等主要跨境市场语言。"),
        ("免费版够用吗？", "每天 3 次适合体验和少量 SKU 测试；大量上新建议开通会员无限次。"),
        ("生成的内容可以直接上传 Amazon 吗？", "建议作为高质量初稿，人工终审后再上传。AI 不替代最终合规责任。"),
        ("手机和电脑都能用吗？", "可以。网页版适合办公，手机版 /mobile/ 支持 PWA 添加到桌面。"),
        ("和 ChatGPT 有什么区别？", "本产品 Prompt 专为 Amazon Listing 优化，内置字符限制、合规检查、批量多语言，无需自己写 Prompt。"),
    ]
    for q, a in faq:
        add_para(doc, f"Q：{q}", bold=True)
        add_para(doc, f"A：{a}")
    doc.add_paragraph()

    add_heading(doc, "十二、推广注意事项（合规）")
    add_bullets(
        doc,
        [
            "宣传时说明「AI 文案仅供参考，上传前请人工终审」",
            "不要承诺「保证爆单」「100% 过审」「官方认证」等无法兑现的话术",
            "竞品对比保持客观，不贬低具体品牌",
            "截图宣传时可打码客户敏感信息",
            "价格以官网实时显示为准",
        ],
    )
    doc.add_paragraph()

    add_heading(doc, "十三、关键信息速查")
    add_bullets(
        doc,
        [
            f"产品名称：跨境 AI Listing 生成器",
            f"官网地址：{SITE_URL}",
            "手机版：https://www.kjdsai.cn/mobile/",
            "支持语言：36 种",
            "免费额度：3 次/天",
            "会员：月卡 ¥29.9 / 半年 ¥69.9 / 年卡 ¥129.9",
            "输出格式：标题 + 5 Bullet + HTML 描述 + Search Terms + Word 导出",
        ],
    )

    doc.save(OUTPUT)
    print(f"已生成: {OUTPUT}")


if __name__ == "__main__":
    build()
