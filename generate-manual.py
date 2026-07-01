# -*- coding: utf-8 -*-
"""生成《跨境 AI Listing 生成器》Word 使用说明书"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

OUTPUT = Path(__file__).parent / "跨境AI Listing生成器-使用说明书.docx"


def set_doc_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        for run in p.runs:
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_highlight_box(doc, title, items):
    add_para(doc, title, bold=True)
    add_bullets(doc, items)
    doc.add_paragraph()


def build():
    doc = Document()
    set_doc_font(doc)

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    add_title(doc, "跨境 AI Listing 生成器")
    add_subtitle(doc, "产品使用说明书")
    doc.add_paragraph()

    add_para(
        doc,
        "本说明书适用于「跨境 AI Listing 生成器」本地版及网页版，帮助跨境卖家快速生成 Amazon 等平台所需的多语言商品 Listing 文案。",
    )
    doc.add_paragraph()

    # 1. 核心优势
    add_heading(doc, "一、核心优势", 1)
    add_highlight_box(
        doc,
        "为什么选择本工具？",
        [
            "一键多语言：支持 36 种语言，覆盖亚洲、欧洲、非洲、美洲四大洲主要跨境市场，不是简单翻译，而是针对各站点做本地化 SEO 优化。",
            "专业 Listing 结构：自动生成标题、核心优势（5 条）、产品描述、关键词，符合 Amazon Listing 规范。",
            "按大洲筛选语言：可按亚洲 / 欧洲 / 非洲 / 美洲筛选，只显示当前大洲语言，避免选项过多、干扰操作。",
            "手动勾选，灵活可控：不强制全选，您可按需勾选任意语言组合，适合不同站点策略。",
            "历史记录本地保存：成功生成后，产品名称、卖点、产品类型会自动保存到浏览器本地，下次可快速填充，无需重复输入。",
            "历史仅保存在本机：历史记录存储在浏览器 localStorage 中，不会上传到任何服务器，保护您的商业信息。",
            "一键复制：生成结果支持分块复制（标题 / 核心优势 / 描述 / 关键词），直接粘贴到卖家后台。",
            "API Key 安全：DeepSeek API Key 仅保存在服务端环境变量中，前端通过本地代理调用，不会暴露在网页代码里。",
            "开箱即用：双击 start.bat 即可启动，无需复杂配置（仅需配置 .env 中的 API Key）。",
            "Powered by DeepSeek：采用 DeepSeek 大模型，生成质量高、响应稳定。",
        ],
    )

    # 2. 环境要求
    add_heading(doc, "二、环境要求", 1)
    add_bullets(
        doc,
        [
            "操作系统：Windows 10 / 11（推荐）",
            "Node.js：LTS 版本（需已安装并加入系统 PATH）",
            "浏览器：Chrome、Edge 等现代浏览器",
            "网络：需能访问 DeepSeek API（api.deepseek.com）",
            "DeepSeek API Key：在 DeepSeek 开放平台申请",
        ],
    )
    doc.add_paragraph()

    # 3. 安装与启动
    add_heading(doc, "三、安装与启动", 1)
    add_heading(doc, "3.1 首次使用", 2)
    add_numbered(
        doc,
        [
            "进入项目文件夹：cross-border-listing-generator",
            "复制 .env.example 为 .env，填入您的 DEEPSEEK_API_KEY",
            "双击运行 start.bat（首次会自动安装依赖，请稍候）",
            "浏览器自动打开 http://127.0.0.1:5173",
            "页面顶部状态条显示绿色「服务器已连接」即可使用",
        ],
    )
    add_heading(doc, "3.2 重要提醒", 2)
    add_bullets(
        doc,
        [
            "请勿直接双击 index.html 打开页面，必须通过 start.bat 或 npm start 启动服务器",
            "运行期间请保持黑色命令行窗口不要关闭，关闭后网页将无法生成 Listing",
            "若端口 5173 被占用，可先执行 taskkill /F /IM node.exe 后重新启动",
        ],
    )
    doc.add_paragraph()

    # 4. 功能说明
    add_heading(doc, "四、功能说明", 1)
    add_heading(doc, "4.1 产品信息填写", 2)
    add_bullets(
        doc,
        [
            "产品名称：输入商品中文名称，如「不锈钢厨房剪刀」",
            "产品卖点 / 特点：填写核心卖点，可多行，越详细生成效果越好",
            "产品类型：如「厨房用品」「宠物用品」等，用于优化 SEO 关键词",
            "文案风格：专业正式 / 活泼口语化，二选一",
        ],
    )

    add_heading(doc, "4.2 历史记录（本地）", 2)
    add_para(doc, "这是本工具的重要便利功能，请务必了解以下说明：")
    add_bullets(
        doc,
        [
            "每次成功生成 Listing 后，系统会自动保存产品名称、卖点、产品类型",
            "在产品名称或产品类型输入框中，可看到历史下拉建议，选中后自动填充全部字段",
            "产品类型下方会显示历史标签（如「产品名 · 类型」），点击即可一键填充",
            "数据仅保存在您当前电脑的当前浏览器中（localStorage）",
            "不会上传到服务器，其他电脑、其他浏览器无法看到您的历史",
            "清除浏览器缓存、换浏览器、重装系统后，历史记录会丢失",
            "最多保存最近 30 条记录",
        ],
    )

    add_heading(doc, "4.3 产品语言选择", 2)
    add_bullets(
        doc,
        [
            "支持 36 种语言，按大洲分类：全部 / 亚洲 / 欧洲 / 非洲 / 美洲",
            "点击某个大洲后，只显示该大洲语言，其他大洲语言隐藏",
            "默认不自动勾选，请手动勾选需要的语言",
            "点击「全部」显示所有语言，同样默认不勾选",
            "可跨大洲勾选：先在「欧洲」勾选德语、法语，再切到「亚洲」勾选日语，生成时会包含所有已勾选语言",
        ],
    )

    add_heading(doc, "4.4 生成结果", 2)
    add_bullets(
        doc,
        [
            "标题：150–200 字符，含 SEO 关键词",
            "核心优势：5 条 bullet points，每条不超过 500 字符",
            "产品描述：200–300 词（日语等按字符计）",
            "关键词：250 字符内，逗号分隔",
            "多语言结果以标签页切换，每块内容均有「一键复制」按钮",
        ],
    )
    doc.add_paragraph()

    # 5. 使用流程
    add_heading(doc, "五、标准使用流程", 1)
    add_numbered(
        doc,
        [
            "双击 start.bat 启动服务",
            "浏览器打开 http://127.0.0.1:5173",
            "填写产品名称、卖点、产品类型（或从历史记录中选择）",
            "选择大洲，勾选需要的产品语言",
            "选择文案风格",
            "点击「生成 Listing」，等待完成（语言越多耗时越长）",
            "在结果区切换语言标签，复制标题、核心优势、描述、关键词到卖家后台",
        ],
    )
    doc.add_paragraph()

    # 6. 部署到 Vercel（可选）
    add_heading(doc, "六、在线部署（可选）", 1)
    add_para(doc, "若需部署到 Vercel 供团队在线使用：")
    add_numbered(
        doc,
        [
            "将项目推送到 GitHub",
            "在 Vercel 导入仓库",
            "在 Environment Variables 中添加 DEEPSEEK_API_KEY",
            "部署完成后通过 Vercel 域名访问",
        ],
    )
    add_para(doc, "注意：在线版的历史记录同样只保存在用户浏览器本地，与部署服务器无关。")
    doc.add_paragraph()

    # 7. 常见问题
    add_heading(doc, "七、常见问题", 1)
    faqs = [
        ("网页打不开？", "请确认已通过 start.bat 启动，并访问 http://127.0.0.1:5173，不要直接打开 HTML 文件。"),
        ("Failed to fetch？", "服务器未运行或命令行窗口已关闭，请重新运行 start.bat。"),
        ("生成超时？", "勾选语言过多时请耐心等待，或减少语言数量后重试。28 种语言约需 5–8 分钟。"),
        ("历史记录不见了？", "历史存在浏览器本地，清除缓存或换浏览器后会丢失，属正常现象。"),
        ("API Key 会泄露吗？", "Key 仅配置在 .env 或 Vercel 环境变量中，不会出现在前端代码或网页源码里。"),
        ("生成的内容可以直接用吗？", "建议生成后人工审核，根据各平台规则微调后再上架。"),
    ]
    for q, a in faqs:
        add_para(doc, f"Q：{q}", bold=True)
        add_para(doc, f"A：{a}")
        doc.add_paragraph()

    # 8. 安全与隐私
    add_heading(doc, "八、安全与隐私说明", 1)
    add_bullets(
        doc,
        [
            "产品信息发送至 DeepSeek API 用于生成文案，请遵守 DeepSeek 服务条款",
            "历史记录仅存储在本地浏览器，不上传、不同步、不共享",
            "请勿将 .env 文件或 API Key 提交到 Git 或分享给他人",
            "若 API Key 曾泄露，请立即在 DeepSeek 平台撤销并重新生成",
        ],
    )
    doc.add_paragraph()

    add_para(doc, "—— 文档结束 ——")
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUTPUT)
    print(f"已生成: {OUTPUT}")


if __name__ == "__main__":
    build()
